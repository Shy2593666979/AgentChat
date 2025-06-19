import os
from docx import Document
from openai import OpenAI
import fitz  # PyMuPDF


class ResumeEnhancer:
    def __init__(self, api_key):
        self.client = OpenAI(api_key=api_key, base_url="")

    def enhance_text(self, text):
        """使用 OpenAI API 优化文本"""
        if not text.strip():  # 如果是空文本就直接返回
            return text

        try:
            response = self.client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system",
                     "content": "你是一个专业的简历优化助手。请优化以下简历内容，使其更专业、更有说服力，但保持事实准确性。保持原有的格式和标点符号。"},
                    {"role": "user", "content": f"请优化以下简历内容，保持相同的格式：\n{text}"}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"调用 API 时出错: {e}")
            return text

    def process_docx(self, input_path, output_path):
        """处理 Word 文档，保留所有格式"""
        doc = Document(input_path)

        # 处理每个段落，保留格式
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 保存原始格式
                original_runs = []
                for run in paragraph.runs:
                    original_runs.append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size,
                        'font_name': run.font.name
                    })

                # 优化文本
                enhanced_text = self.enhance_text(paragraph.text)

                # 清除原有内容
                paragraph.clear()

                # 如果只有一个run，直接设置增强后的文本
                if len(original_runs) == 1:
                    run = paragraph.add_run(enhanced_text)
                    # 恢复格式
                    format_info = original_runs[0]
                    run.bold = format_info['bold']
                    run.italic = format_info['italic']
                    run.underline = format_info['underline']
                    run.font.size = format_info['font_size']
                    run.font.name = format_info['font_name']
                else:
                    # 尝试保持原有的文本分段
                    words = enhanced_text.split()
                    total_len = len(words)
                    runs_len = len(original_runs)
                    words_per_run = total_len // runs_len

                    for i, format_info in enumerate(original_runs):
                        # 计算当前run应该包含的文本
                        start_idx = i * words_per_run
                        end_idx = (i + 1) * words_per_run if i < runs_len - 1 else total_len
                        text_segment = ' '.join(words[start_idx:end_idx])

                        # 添加新的run并应用格式
                        run = paragraph.add_run(text_segment + ' ')
                        run.bold = format_info['bold']
                        run.italic = format_info['italic']
                        run.underline = format_info['underline']
                        run.font.size = format_info['font_size']
                        run.font.name = format_info['font_name']

        # 处理表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            # 保存原始格式
                            original_runs = []
                            for run in paragraph.runs:
                                original_runs.append({
                                    'text': run.text,
                                    'bold': run.bold,
                                    'italic': run.italic,
                                    'underline': run.underline,
                                    'font_size': run.font.size,
                                    'font_name': run.font.name
                                })

                            # 优化文本
                            enhanced_text = self.enhance_text(paragraph.text)

                            # 清除原有内容
                            paragraph.clear()

                            # 应用相同的格式处理逻辑
                            if len(original_runs) == 1:
                                run = paragraph.add_run(enhanced_text)
                                format_info = original_runs[0]
                                run.bold = format_info['bold']
                                run.italic = format_info['italic']
                                run.underline = format_info['underline']
                                run.font.size = format_info['font_size']
                                run.font.name = format_info['font_name']

        # 保存修改后的文档
        doc.save(output_path)

    def process_pdf(self, input_path, output_path):
        """处理 PDF 文档"""
        # 打开PDF文件
        doc = fitz.open(input_path)
        new_doc = fitz.open()  # 创建新文档

        for page_num in range(len(doc)):
            page = doc[page_num]
            new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)

            # 获取页面上的各个文本块
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            # 获取原始格式信息
                            original_font = span["font"]
                            original_size = span["size"]
                            original_color = span["color"]
                            original_text = span["text"]
                            original_bbox = span["bbox"]

                            # 优化文本内容
                            if original_text.strip():
                                enhanced_text = self.enhance_text(original_text)
                            else:
                                enhanced_text = original_text

                            # 使用原始格式插入增强后的文本
                            new_page.insert_text(
                                point=(original_bbox[0], original_bbox[1]),
                                text=enhanced_text,
                                fontname=original_font,
                                fontsize=original_size,
                                color=original_color
                            )

        # 保存新文档
        new_doc.save(output_path)
        new_doc.close()
        doc.close()


def main():
    # 设置你的 OpenAI API 密钥
    api_key = ""
    enhancer = ResumeEnhancer(api_key)

    # 输入文件路径
    input_file = "input_resume.docx"  # 或 "input_resume.pdf"

    # 输出文件路径
    output_file = "enhanced_resume" + os.path.splitext(input_file)[1]

    # 根据文件类型处理
    if input_file.endswith('.docx'):
        enhancer.process_docx(input_file, output_file)
    elif input_file.endswith('.pdf'):
        enhancer.process_pdf(input_file, output_file)
    else:
        print("不支持的文件格式")

    print(f"简历优化完成，已保存至: {output_file}")


if __name__ == "__main__":
    main()


